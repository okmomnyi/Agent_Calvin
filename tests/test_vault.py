"""Study Vault tests: embeddings/cosine, extraction+chunking, idempotent ingest,
cited answers, and the 'not in your notes' guardrail (never guess from thin air)."""

from __future__ import annotations

import pytest

from core.embeddings import HashingEmbedder, cosine, pack_vector, unpack_vector
from core.doc_extract import Passage, chunk_passages, extract
from core.llm import LLMClient
from skills.study_vault import VaultSkill


# ------------------------------------------------------------------ embeddings
def test_hashing_embedder_deterministic_and_stable():
    e = HashingEmbedder(dim=256)
    v1 = e.embed("binary search tree logarithmic lookup")
    v2 = e.embed("binary search tree logarithmic lookup")
    assert v1 == v2                      # deterministic (stable hash, not salted)
    assert len(v1) == 256


def test_cosine_self_is_one_and_related_higher():
    e = HashingEmbedder()
    a = e.embed("photosynthesis converts light energy in chloroplasts")
    b = e.embed("chloroplasts perform photosynthesis using light")
    c = e.embed("quantum entanglement teleportation qubits")
    assert cosine(a, a) == pytest.approx(1.0, abs=1e-6)
    assert cosine(a, b) > cosine(a, c)   # topical overlap beats unrelated


def test_pack_unpack_roundtrip():
    v = [0.1, -0.2, 0.3, 0.0]
    assert unpack_vector(pack_vector(v)) == pytest.approx(v, abs=1e-6)


# ------------------------------------------------------------------ extraction + chunking
def test_extract_txt_and_md(tmp_path):
    (tmp_path / "a.txt").write_text("hello world notes", encoding="utf-8")
    (tmp_path / "b.md").write_text("# Heading\nsome markdown", encoding="utf-8")
    assert extract(tmp_path / "a.txt")[0].text == "hello world notes"
    assert "markdown" in extract(tmp_path / "b.md")[0].text


def test_extract_docx(tmp_path):
    import docx

    d = docx.Document()
    d.add_paragraph("Automata theory covers finite state machines.")
    d.add_paragraph("A DFA has no epsilon transitions.")
    path = tmp_path / "notes.docx"
    d.save(str(path))
    text = extract(path)[0].text
    assert "finite state machines" in text and "DFA" in text


def test_chunk_passages_splits_with_overlap():
    words = " ".join(f"w{i}" for i in range(1200))
    chunks = chunk_passages([Passage("p.1", words)], size_words=500, overlap=60)
    assert len(chunks) >= 3
    assert all(c.loc == "p.1" for c in chunks)


def test_unsupported_file_returns_empty(tmp_path):
    (tmp_path / "x.xyz").write_text("nope", encoding="utf-8")
    assert extract(tmp_path / "x.xyz") == []


# ------------------------------------------------------------------ vault skill
class _VaultLLM(LLMClient):
    def __init__(self, answer="A BST gives O(log n) lookup because the tree is ordered [1]."):
        self.routes = {"default": "m", "research": "m"}
        self.defaults = {}
        self._answer = answer

    def chat(self, task, messages, **kw):  # type: ignore[override]
        return self._answer


@pytest.fixture
def vault(mem, tmp_path, monkeypatch):
    import skills.study_vault as sv

    real = sv.get_settings()

    class _S:
        def __init__(self): self.data_dir = tmp_path
        def __getattr__(self, n): return getattr(real, n)

    monkeypatch.setattr(sv, "get_settings", lambda: _S())
    skill = VaultSkill(memory=mem, embedder=HashingEmbedder(), llm=_VaultLLM())
    # seed a unit with a note
    unit = tmp_path / "vault" / "CS301"
    unit.mkdir(parents=True)
    (unit / "trees.txt").write_text(
        "Binary search trees allow logarithmic lookup, insertion and deletion. "
        "A balanced BST keeps its height proportional to log n, so operations stay O(log n).",
        encoding="utf-8")
    return skill, tmp_path


def test_ingest_is_idempotent(vault):
    skill, _ = vault
    first = skill.ingest()
    assert first.data["ingested"] == 1 and first.data["chunks"] >= 1
    second = skill.ingest()
    assert second.data["ingested"] == 0 and second.data["skipped"] == 1  # unchanged -> skipped


def test_ask_answers_with_citation(vault):
    skill, _ = vault
    skill.ingest()
    result = skill.ask(question="why is lookup in a binary search tree logarithmic?")
    assert "From your notes:" in result.text
    assert "trees.txt" in result.text
    assert result.data["best_score"] >= skill.min_score


def test_ask_not_in_notes_when_unrelated(vault):
    skill, _ = vault
    skill.ingest()
    result = skill.ask(question="explain quantum chromodynamics gluon confinement",
                       allow_web=False)
    assert "not in your notes" in result.text.lower()
    assert result.data["in_notes"] is False


def test_status_lists_units(vault):
    skill, _ = vault
    skill.ingest()
    st = skill.status()
    assert "CS301" in st.text
    assert st.data["total_chunks"] >= 1


def test_retrieve_returns_file_and_loc(vault):
    skill, _ = vault
    skill.ingest()
    hits = skill.retrieve("balanced BST height log n")
    assert hits and hits[0]["file"] == "trees.txt"
    assert hits[0]["unit"] == "CS301"
